from django.db import models


def extract_docx_text(file_obj) -> str:
    """Extract plain text from an uploaded .docx file."""
    from docx import Document as DocxDocument
    file_obj.seek(0)
    doc = DocxDocument(file_obj)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


class Document(models.Model):
    title = models.CharField(max_length=255)
    file = models.FileField(upload_to="documents/", blank=True, null=True)
    content = models.TextField(blank=True, help_text="Extracted full text.")
    chunk_count = models.PositiveIntegerField(default=0)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        ordering = ["-created_at"]

    def __str__(self) -> str:
        return self.title

    def reindex(self) -> int:
        from qa.rag import index_document
        self.chunk_count = index_document(self.id, self.title, self.content or "")
        Document.objects.filter(pk=self.pk).update(chunk_count=self.chunk_count)
        return self.chunk_count
