"""Generate a sample DOCX for testing."""
from docx import Document

doc = Document()
doc.add_heading("LangChain Overview", level=1)
doc.add_paragraph(
    "LangChain is an open-source framework for building applications powered by "
    "large language models (LLMs). It provides composable abstractions for "
    "prompts, chains, retrieval, memory, agents, and tool use."
)
doc.add_heading("Retrieval-Augmented Generation", level=2)
doc.add_paragraph(
    "Retrieval-Augmented Generation (RAG) combines a retriever, which fetches "
    "relevant documents from a vector store, with a generator LLM that produces "
    "an answer grounded in those documents. RAG reduces hallucinations and lets "
    "models answer questions about private or up-to-date data without retraining."
)
doc.add_heading("Vector Stores", level=2)
doc.add_paragraph(
    "A vector store indexes embeddings of text chunks so that similar pieces can "
    "be retrieved via approximate nearest neighbor search. Common choices include "
    "Chroma, FAISS, Pinecone, Weaviate, and pgvector."
)
doc.save("sample_data/langchain_overview.docx")

doc2 = Document()
doc2.add_heading("Django in a Nutshell", level=1)
doc2.add_paragraph(
    "Django is a high-level Python web framework that encourages rapid "
    "development and clean, pragmatic design. It ships with an ORM, an "
    "auto-generated admin interface, a templating engine, and a robust auth "
    "system."
)
doc2.add_heading("The Admin Site", level=2)
doc2.add_paragraph(
    "Django Admin reads model metadata to provide a quick interface for "
    "creating, reading, updating, and deleting database records. It is the "
    "primary user interface used by this project."
)
doc2.save("sample_data/django_overview.docx")
print("Wrote 2 sample DOCX files.")
